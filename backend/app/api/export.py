"""Export rewritten bullets back to a DOCX file."""
from __future__ import annotations

import io

from fastapi import APIRouter
from fastapi.responses import StreamingResponse

from app.schemas import ExportRequest

router = APIRouter()


@router.post("/export")
async def export_docx(req: ExportRequest) -> StreamingResponse:
    from docx import Document

    doc = Document()
    doc.add_heading(req.name, level=1)
    doc.add_heading("Experience", level=2)
    for bullet in req.bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{req.name.replace(' ', '_')}.docx"
    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
